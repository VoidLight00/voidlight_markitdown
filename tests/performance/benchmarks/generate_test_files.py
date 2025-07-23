#!/usr/bin/env python3
"""
Generate large test files for performance benchmarking.

This script creates test files of various sizes and formats to test
the performance of voidlight_markitdown with large files.
"""

import os
import random
import string
import json
import struct
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional
import lorem
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import HexColor
import pandas as pd
import numpy as np

class TestFileGenerator:
    """Generate various types of test files for performance testing."""
    
    def __init__(self, base_path: str = "/Users/voidlight/voidlight_markitdown/packages/voidlight_markitdown/tests/performance/test_files"):
        self.base_path = Path(base_path)
        self.base_path.mkdir(parents=True, exist_ok=True)
        
        # Korean text samples
        self.korean_samples = [
            "안녕하세요. 이것은 한국어 테스트 문장입니다.",
            "성능 테스트를 위한 대용량 파일을 생성하고 있습니다.",
            "마크다운 변환 도구의 처리 속도를 측정합니다.",
            "다양한 형식의 문서를 지원합니다.",
            "스트림 기반 처리로 메모리 효율성을 높입니다.",
            "한글 문서의 특수한 처리 요구사항을 고려합니다.",
            "인공지능과 머신러닝 기술을 활용합니다.",
            "빅데이터 처리와 분석 능력을 향상시킵니다.",
            "클라우드 컴퓨팅 환경에서의 성능을 최적화합니다.",
            "사용자 경험을 개선하기 위해 지속적으로 노력합니다."
        ]
        
        # Mixed content samples
        self.mixed_samples = [
            "This is a test sentence in English. 이것은 한국어 테스트 문장입니다.",
            "Performance testing with 성능 테스트 large files 대용량 파일.",
            "Machine Learning 머신러닝 and AI 인공지능 technologies.",
            "Cloud Computing 클라우드 컴퓨팅 optimization 최적화.",
            "Data Science 데이터 과학 and Analytics 분석.",
        ]
        
        print(f"Test files will be generated in: {self.base_path}")
    
    def generate_text_file(self, size_mb: int, language: str = "english") -> str:
        """Generate a plain text file of specified size."""
        filename = f"test_{language}_{size_mb}mb.txt"
        filepath = self.base_path / filename
        
        print(f"Generating {filename}...")
        
        target_size = size_mb * 1024 * 1024  # Convert to bytes
        current_size = 0
        
        with open(filepath, 'w', encoding='utf-8') as f:
            # Write header
            header = f"# Large Text File - {size_mb}MB - {language.title()}\n"
            header += f"Generated: {datetime.now().isoformat()}\n"
            header += f"Purpose: Performance testing for voidlight_markitdown\n"
            header += "=" * 80 + "\n\n"
            f.write(header)
            current_size += len(header.encode('utf-8'))
            
            paragraph_count = 0
            while current_size < target_size:
                paragraph_count += 1
                
                if language == "english":
                    # Generate English paragraph
                    paragraph = f"\n## Section {paragraph_count}\n\n"
                    paragraph += lorem.paragraph() + "\n\n"
                    
                    # Add some structured data occasionally
                    if paragraph_count % 10 == 0:
                        paragraph += self._generate_table_text() + "\n\n"
                    
                elif language == "korean":
                    # Generate Korean paragraph
                    paragraph = f"\n## 섹션 {paragraph_count}\n\n"
                    # Combine random Korean sentences
                    num_sentences = random.randint(5, 10)
                    for _ in range(num_sentences):
                        paragraph += random.choice(self.korean_samples) + " "
                    paragraph += "\n\n"
                    
                    # Add Korean table occasionally
                    if paragraph_count % 10 == 0:
                        paragraph += self._generate_korean_table_text() + "\n\n"
                    
                elif language == "mixed":
                    # Generate mixed content
                    paragraph = f"\n## Section 섹션 {paragraph_count}\n\n"
                    # Mix English and Korean
                    num_sentences = random.randint(5, 10)
                    for i in range(num_sentences):
                        if i % 2 == 0:
                            paragraph += lorem.sentence() + " "
                        else:
                            paragraph += random.choice(self.korean_samples) + " "
                    paragraph += "\n\n"
                    
                    # Add mixed table occasionally
                    if paragraph_count % 10 == 0:
                        paragraph += self._generate_mixed_table_text() + "\n\n"
                
                f.write(paragraph)
                current_size += len(paragraph.encode('utf-8'))
                
                # Progress indicator
                if paragraph_count % 100 == 0:
                    progress = (current_size / target_size) * 100
                    print(f"  Progress: {progress:.1f}%")
        
        actual_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        print(f"  Generated: {filename} ({actual_size_mb:.2f}MB)")
        return str(filepath)
    
    def _generate_table_text(self) -> str:
        """Generate a text table."""
        table = "| Column A | Column B | Column C | Value |\n"
        table += "|----------|----------|----------|-------|\n"
        for i in range(10):
            table += f"| Item {i+1} | Data {i+1} | Info {i+1} | {random.randint(100, 999)} |\n"
        return table
    
    def _generate_korean_table_text(self) -> str:
        """Generate a Korean text table."""
        table = "| 항목 | 설명 | 값 | 상태 |\n"
        table += "|------|------|-----|------|\n"
        items = ["첫번째", "두번째", "세번째", "네번째", "다섯번째", "여섯번째", "일곱번째", "여덟번째", "아홉번째", "열번째"]
        statuses = ["완료", "진행중", "대기", "취소"]
        for i, item in enumerate(items):
            table += f"| {item} | 데이터 {i+1} | {random.randint(100, 999)} | {random.choice(statuses)} |\n"
        return table
    
    def _generate_mixed_table_text(self) -> str:
        """Generate a mixed language table."""
        table = "| English | 한국어 | Value 값 | Status 상태 |\n"
        table += "|---------|--------|----------|-------------|\n"
        for i in range(10):
            eng_word = lorem.sentence().split()[0]
            kor_word = random.choice(self.korean_samples).split()[0]
            status = random.choice(["Complete 완료", "In Progress 진행중", "Pending 대기"])
            table += f"| {eng_word} | {kor_word} | {random.randint(100, 999)} | {status} |\n"
        return table
    
    def generate_pdf_file(self, size_mb: int, include_images: bool = True, language: str = "english") -> str:
        """Generate a PDF file of specified size."""
        filename = f"test_pdf_{language}_{size_mb}mb.pdf"
        filepath = self.base_path / filename
        
        print(f"Generating {filename}...")
        
        # Create PDF
        c = canvas.Canvas(str(filepath), pagesize=letter)
        width, height = letter
        
        # Try to register Korean font if available
        try:
            # Common Korean font paths on different systems
            korean_fonts = [
                "/System/Library/Fonts/AppleSDGothicNeo.ttc",  # macOS
                "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",  # Linux
                "C:\\Windows\\Fonts\\malgun.ttf",  # Windows
            ]
            
            font_registered = False
            for font_path in korean_fonts:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('Korean', font_path))
                        font_registered = True
                        break
                    except:
                        pass
            
            if not font_registered and language in ["korean", "mixed"]:
                print("  Warning: Korean font not found, using default font")
        except:
            pass
        
        # Generate content
        target_size = size_mb * 1024 * 1024
        current_size = 0
        page_num = 0
        y_position = height - inch
        
        while current_size < target_size:
            page_num += 1
            
            # Add page header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(inch, height - 0.5*inch, f"Page {page_num}")
            y_position = height - inch
            
            # Add content based on language
            c.setFont("Helvetica", 12)
            
            while y_position > inch:
                if language == "english":
                    text = lorem.paragraph()
                elif language == "korean":
                    # Use Korean text if font is available
                    text = " ".join(random.choice(self.korean_samples) for _ in range(3))
                    try:
                        c.setFont("Korean", 12)
                    except:
                        # Fallback to English if Korean font not available
                        text = f"[Korean Text: {len(text)} characters]"
                        c.setFont("Helvetica", 12)
                else:  # mixed
                    text = random.choice(self.mixed_samples) * 3
                
                # Wrap text
                lines = self._wrap_text(text, 80)
                for line in lines:
                    if y_position < inch:
                        break
                    c.drawString(inch, y_position, line)
                    y_position -= 14
                
                # Add image occasionally if requested
                if include_images and random.random() < 0.1 and y_position > 3*inch:
                    # Draw a simple rectangle as "image"
                    c.setFillColor(HexColor('#EEEEEE'))
                    c.rect(inch, y_position - 2*inch, 4*inch, 2*inch, fill=True)
                    c.setFillColor(HexColor('#000000'))
                    c.drawString(2.5*inch, y_position - inch, "Sample Image")
                    y_position -= 2.5*inch
            
            c.showPage()
            
            # Estimate current size (rough approximation)
            current_size = page_num * 50000  # Approximate bytes per page
            
            if page_num % 50 == 0:
                progress = (current_size / target_size) * 100
                print(f"  Progress: {progress:.1f}% ({page_num} pages)")
        
        c.save()
        
        actual_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        print(f"  Generated: {filename} ({actual_size_mb:.2f}MB, {page_num} pages)")
        return str(filepath)
    
    def _wrap_text(self, text: str, max_chars: int) -> List[str]:
        """Simple text wrapping."""
        words = text.split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 > max_chars:
                lines.append(" ".join(current_line))
                current_line = [word]
                current_length = len(word)
            else:
                current_line.append(word)
                current_length += len(word) + 1
        
        if current_line:
            lines.append(" ".join(current_line))
        
        return lines
    
    def generate_docx_file(self, size_mb: int, language: str = "english") -> str:
        """Generate a DOCX file of specified size."""
        filename = f"test_docx_{language}_{size_mb}mb.docx"
        filepath = self.base_path / filename
        
        print(f"Generating {filename}...")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Large DOCX Test File - {size_mb}MB', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f'Generated: {datetime.now().isoformat()}')
        doc.add_paragraph(f'Language: {language.title()}')
        doc.add_paragraph('Purpose: Performance testing for voidlight_markitdown')
        
        # Generate content
        target_size = size_mb * 1024 * 1024
        current_size = 0
        section_num = 0
        
        while current_size < target_size:
            section_num += 1
            
            # Add section heading
            if language == "english":
                doc.add_heading(f'Section {section_num}', level=1)
            elif language == "korean":
                doc.add_heading(f'섹션 {section_num}', level=1)
            else:
                doc.add_heading(f'Section 섹션 {section_num}', level=1)
            
            # Add paragraphs
            for _ in range(random.randint(5, 10)):
                p = doc.add_paragraph()
                
                if language == "english":
                    p.add_run(lorem.paragraph())
                elif language == "korean":
                    text = " ".join(random.choice(self.korean_samples) for _ in range(5))
                    p.add_run(text)
                else:
                    # Mixed content
                    for i in range(5):
                        if i % 2 == 0:
                            p.add_run(lorem.sentence() + " ")
                        else:
                            p.add_run(random.choice(self.korean_samples) + " ")
            
            # Add table occasionally
            if section_num % 5 == 0:
                self._add_docx_table(doc, language)
            
            # Add list occasionally
            if section_num % 7 == 0:
                self._add_docx_list(doc, language)
            
            # Save and check size
            doc.save(filepath)
            current_size = os.path.getsize(filepath)
            
            if section_num % 10 == 0:
                progress = (current_size / target_size) * 100
                print(f"  Progress: {progress:.1f}% ({section_num} sections)")
        
        actual_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        print(f"  Generated: {filename} ({actual_size_mb:.2f}MB)")
        return str(filepath)
    
    def _add_docx_table(self, doc: Document, language: str):
        """Add a table to DOCX document."""
        if language == "english":
            headers = ['Item', 'Description', 'Value', 'Status']
        elif language == "korean":
            headers = ['항목', '설명', '값', '상태']
        else:
            headers = ['Item 항목', 'Description 설명', 'Value 값', 'Status 상태']
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        # Add data rows
        for i in range(10):
            row_cells = table.add_row().cells
            if language == "english":
                row_cells[0].text = f'Item {i+1}'
                row_cells[1].text = lorem.sentence()
                row_cells[2].text = str(random.randint(100, 999))
                row_cells[3].text = random.choice(['Active', 'Pending', 'Complete'])
            elif language == "korean":
                row_cells[0].text = f'항목 {i+1}'
                row_cells[1].text = random.choice(self.korean_samples)
                row_cells[2].text = str(random.randint(100, 999))
                row_cells[3].text = random.choice(['활성', '대기', '완료'])
            else:
                row_cells[0].text = f'Item 항목 {i+1}'
                row_cells[1].text = random.choice(self.mixed_samples)
                row_cells[2].text = str(random.randint(100, 999))
                row_cells[3].text = random.choice(['Active 활성', 'Pending 대기', 'Complete 완료'])
        
        doc.add_paragraph()  # Add spacing
    
    def _add_docx_list(self, doc: Document, language: str):
        """Add a list to DOCX document."""
        if language == "english":
            doc.add_paragraph('Key Features:', style='Heading 2')
            items = [lorem.sentence() for _ in range(5)]
        elif language == "korean":
            doc.add_paragraph('주요 기능:', style='Heading 2')
            items = random.sample(self.korean_samples, 5)
        else:
            doc.add_paragraph('Key Features 주요 기능:', style='Heading 2')
            items = random.sample(self.mixed_samples, 5)
        
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_paragraph()  # Add spacing
    
    def generate_excel_file(self, size_mb: int, language: str = "english") -> str:
        """Generate an Excel file of specified size."""
        filename = f"test_excel_{language}_{size_mb}mb.xlsx"
        filepath = self.base_path / filename
        
        print(f"Generating {filename}...")
        
        wb = openpyxl.Workbook()
        
        # Generate multiple sheets
        target_size = size_mb * 1024 * 1024
        sheet_num = 0
        
        while True:
            sheet_num += 1
            
            if sheet_num == 1:
                ws = wb.active
                ws.title = f"Sheet{sheet_num}"
            else:
                ws = wb.create_sheet(f"Sheet{sheet_num}")
            
            # Add headers
            if language == "english":
                headers = ['ID', 'Name', 'Description', 'Category', 'Value', 'Date', 'Status', 'Notes']
            elif language == "korean":
                headers = ['아이디', '이름', '설명', '카테고리', '값', '날짜', '상태', '비고']
            else:
                headers = ['ID', 'Name 이름', 'Description 설명', 'Category 카테고리', 
                          'Value 값', 'Date 날짜', 'Status 상태', 'Notes 비고']
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            
            # Add data rows
            row_count = min(10000, size_mb * 1000)  # Scale rows with file size
            
            for row in range(2, row_count + 2):
                if language == "english":
                    ws.cell(row=row, column=1, value=row-1)
                    ws.cell(row=row, column=2, value=f"Item {row-1}")
                    ws.cell(row=row, column=3, value=lorem.sentence())
                    ws.cell(row=row, column=4, value=random.choice(['A', 'B', 'C', 'D']))
                    ws.cell(row=row, column=5, value=random.randint(100, 9999))
                    ws.cell(row=row, column=6, value=datetime.now().date())
                    ws.cell(row=row, column=7, value=random.choice(['Active', 'Inactive', 'Pending']))
                    ws.cell(row=row, column=8, value=lorem.sentence())
                elif language == "korean":
                    ws.cell(row=row, column=1, value=row-1)
                    ws.cell(row=row, column=2, value=f"항목 {row-1}")
                    ws.cell(row=row, column=3, value=random.choice(self.korean_samples))
                    ws.cell(row=row, column=4, value=random.choice(['가', '나', '다', '라']))
                    ws.cell(row=row, column=5, value=random.randint(100, 9999))
                    ws.cell(row=row, column=6, value=datetime.now().date())
                    ws.cell(row=row, column=7, value=random.choice(['활성', '비활성', '대기']))
                    ws.cell(row=row, column=8, value=random.choice(self.korean_samples))
                else:
                    ws.cell(row=row, column=1, value=row-1)
                    ws.cell(row=row, column=2, value=f"Item 항목 {row-1}")
                    ws.cell(row=row, column=3, value=random.choice(self.mixed_samples))
                    ws.cell(row=row, column=4, value=random.choice(['A 가', 'B 나', 'C 다', 'D 라']))
                    ws.cell(row=row, column=5, value=random.randint(100, 9999))
                    ws.cell(row=row, column=6, value=datetime.now().date())
                    ws.cell(row=row, column=7, value=random.choice(['Active 활성', 'Inactive 비활성', 'Pending 대기']))
                    ws.cell(row=row, column=8, value=random.choice(self.mixed_samples))
                
                if row % 1000 == 0:
                    print(f"  Sheet {sheet_num}: {row} rows")
            
            # Add formulas to last row
            formula_row = row_count + 3
            ws.cell(row=formula_row, column=4, value="Total:")
            ws.cell(row=formula_row, column=5, value=f"=SUM(E2:E{row_count+1})")
            
            # Save and check size
            wb.save(filepath)
            current_size = os.path.getsize(filepath)
            
            if current_size >= target_size:
                break
            
            print(f"  Progress: {(current_size / target_size) * 100:.1f}% ({sheet_num} sheets)")
        
        actual_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        print(f"  Generated: {filename} ({actual_size_mb:.2f}MB, {sheet_num} sheets)")
        return str(filepath)
    
    def generate_all_test_files(self) -> Dict[str, List[str]]:
        """Generate all test files for benchmarking."""
        print("Generating comprehensive test file suite...")
        print("=" * 60)
        
        results = {
            'text': [],
            'pdf': [],
            'docx': [],
            'excel': [],
        }
        
        # File sizes to generate (in MB)
        sizes = [10, 50, 100, 500]
        languages = ['english', 'korean', 'mixed']
        
        # Generate text files
        print("\n1. Generating text files...")
        for size in sizes:
            for lang in languages:
                filepath = self.generate_text_file(size, lang)
                results['text'].append(filepath)
        
        # Generate PDF files
        print("\n2. Generating PDF files...")
        for size in sizes[:3]:  # Skip 500MB for PDFs
            for lang in languages:
                filepath = self.generate_pdf_file(size, include_images=True, language=lang)
                results['pdf'].append(filepath)
        
        # Generate DOCX files
        print("\n3. Generating DOCX files...")
        for size in sizes[:3]:  # Skip 500MB for DOCX
            for lang in languages:
                filepath = self.generate_docx_file(size, lang)
                results['docx'].append(filepath)
        
        # Generate Excel files
        print("\n4. Generating Excel files...")
        for size in [10, 50]:  # Only smaller sizes for Excel
            for lang in languages:
                filepath = self.generate_excel_file(size, lang)
                results['excel'].append(filepath)
        
        # Save file manifest
        manifest_path = self.base_path / "test_files_manifest.json"
        with open(manifest_path, 'w', encoding='utf-8') as f:
            json.dump({
                'generated': datetime.now().isoformat(),
                'base_path': str(self.base_path),
                'files': results,
                'total_files': sum(len(v) for v in results.values()),
            }, f, indent=2)
        
        print(f"\n✅ Test file generation complete!")
        print(f"📁 Files saved to: {self.base_path}")
        print(f"📋 Manifest: {manifest_path}")
        print(f"📊 Total files generated: {sum(len(v) for v in results.values())}")
        
        return results


def main():
    """Main entry point."""
    generator = TestFileGenerator()
    results = generator.generate_all_test_files()
    
    # Print summary
    print("\n" + "=" * 60)
    print("GENERATION SUMMARY")
    print("=" * 60)
    
    for file_type, files in results.items():
        print(f"\n{file_type.upper()} Files ({len(files)}):")
        for filepath in sorted(files):
            size_mb = os.path.getsize(filepath) / (1024 * 1024)
            print(f"  - {os.path.basename(filepath)} ({size_mb:.2f}MB)")


if __name__ == "__main__":
    main()