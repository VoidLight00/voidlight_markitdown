#!/usr/bin/env python3
from voidlight_markitdown import VoidLightMarkItDown
import tempfile
import os

# Test PDF conversion
print('Testing PDF converter...')
try:
    md = VoidLightMarkItDown()
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
        # Create minimal PDF content
        f.write(b'%PDF-1.4\n1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n3 0 obj\n<< /Type /Page /Parent 2 0 R /Resources << >> /MediaBox [0 0 612 792] >>\nendobj\nxref\n0 4\n0000000000 65535 f\n0000000009 00000 n\n0000000058 00000 n\n0000000115 00000 n\ntrailer\n<< /Size 4 /Root 1 0 R >>\nstartxref\n208\n%%EOF')
        f.flush()
        result = md.convert(f.name)
        print(f'PDF conversion result: {result.markdown[:50]}...')
        os.unlink(f.name)
    print('✓ PDF converter available')
except Exception as e:
    print(f'✗ PDF converter error: {e}')

# Test DOCX conversion
print('\nTesting DOCX converter...')
try:
    # Create a simple DOCX file using python-docx
    from docx import Document
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        result = md.convert(f.name)
        print(f'DOCX conversion result: {result.markdown[:50]}...')
        os.unlink(f.name)
    print('✓ DOCX converter available')
except Exception as e:
    print(f'✗ DOCX converter error: {e}')

# Test Excel conversion  
print('\nTesting Excel converter...')
try:
    # Create a simple Excel file using openpyxl
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws['A1'] = 'Hello'
    ws['B1'] = 'World'
    with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as f:
        wb.save(f.name)
        result = md.convert(f.name)
        print(f'Excel conversion result: {result.markdown[:50]}...')
        os.unlink(f.name)
    print('✓ Excel converter available')
except Exception as e:
    print(f'✗ Excel converter error: {e}')

# Test Image/OCR conversion
print('\nTesting Image/OCR converter...')
try:
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
        # Create minimal PNG
        f.write(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\xf0\xb7N\xa1\x00\x00\x00\x00IEND\xaeB`\x82')
        f.flush()
        result = md.convert(f.name)
        print(f'Image conversion result: {result.markdown[:50]}...')
        os.unlink(f.name)
    print('✓ Image converter available')
except Exception as e:
    print(f'✗ Image converter error: {e}')

# Test Korean mode
print('\nTesting Korean mode...')
try:
    md_korean = VoidLightMarkItDown(korean_mode=True)
    # Save Korean text to file first
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
        f.write('안녕하세요. 이것은 한국어 테스트입니다.')
        f.flush()
        result = md_korean.convert(f.name)
        print(f'Korean text result: {result.markdown}')
        os.unlink(f.name)
    print('✓ Korean mode available')
except Exception as e:
    print(f'✗ Korean mode error: {e}')

# List all available converters
print('\n\nListing all converters:')
# md._converters is a list of converters, not a dict
for converter in md._converters:
    print(f'- {type(converter).__name__}')